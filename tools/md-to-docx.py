#!/usr/bin/env python3
"""Convert a FounderHive markdown draft to a clean, readable .docx.

Usage:  python3 tools/md-to-docx.py <input.md> [output.docx]
Handles: # H1 title, ## / ### subheads, blank-line paragraphs, **bold**, *italic*,
and - bullet lists. Requires python-docx (pip install python-docx).
"""
import re, sys, os
from docx import Document
from docx.shared import Pt, RGBColor

def add_runs(p, s):
    for t in re.split(r"(\*\*.+?\*\*|\*.+?\*)", s):
        if not t:
            continue
        if t.startswith("**") and t.endswith("**"):
            p.add_run(t[2:-2]).bold = True
        elif t.startswith("*") and t.endswith("*"):
            p.add_run(t[1:-1]).italic = True
        else:
            p.add_run(t)

def convert(src, out=None):
    out = out or os.path.splitext(src)[0] + ".docx"
    text = open(src, encoding="utf-8").read()
    doc = Document()
    base = doc.styles["Normal"]; base.font.name = "Georgia"; base.font.size = Pt(11)
    for block in (b.strip() for b in re.split(r"\n\s*\n", text) if b.strip()):
        line = block.replace("\n", " ").strip()
        if block.startswith("# "):
            doc.add_heading(level=0).add_run(block[2:].strip()).font.color.rgb = RGBColor(0x1a, 0x1d, 0x24)
        elif block.startswith("### "):
            doc.add_heading(block[4:].strip(), level=2)
        elif block.startswith("## "):
            doc.add_heading(block[3:].strip(), level=1)
        elif block.startswith(("- ", "* ")):
            for li in block.split("\n"):
                p = doc.add_paragraph(style="List Bullet"); add_runs(p, li[2:].strip())
        else:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(10); p.paragraph_format.line_spacing = 1.25
            add_runs(p, line)
    doc.save(out)
    return out

if __name__ == "__main__":
    if len(sys.argv) < 2:
        sys.exit("Usage: python3 tools/md-to-docx.py <input.md> [output.docx]")
    print("wrote", convert(sys.argv[1], sys.argv[2] if len(sys.argv) > 2 else None))
